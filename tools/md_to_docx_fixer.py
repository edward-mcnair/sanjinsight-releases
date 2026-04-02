#!/usr/bin/env python3
"""
md_to_docx_fixer.py — Convert raw-markdown .docx files to properly formatted Word.

Reads a .docx whose paragraphs contain Markdown text (headings, bold,
bullets, tables, code blocks, blockquotes, links, etc.) and rewrites
them using native Word styles and formatting.

Usage:
    python3 tools/md_to_docx_fixer.py input.docx output.docx
"""

import re
import sys
from copy import deepcopy
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ── Microsanj brand colors ──────────────────────────────────────────
BRAND_BLUE     = RGBColor(0x00, 0x6C, 0xBF)
HEADING_COLOR  = RGBColor(0x1A, 0x1A, 0x2E)
CODE_BG        = RGBColor(0xF5, 0xF5, 0xF5)
CODE_FG        = RGBColor(0x33, 0x33, 0x33)
TABLE_HEADER   = RGBColor(0x00, 0x6C, 0xBF)
TABLE_ALT_ROW  = RGBColor(0xF2, 0xF6, 0xFA)
BLOCKQUOTE_BAR = RGBColor(0x00, 0x6C, 0xBF)
BLOCKQUOTE_BG  = RGBColor(0xF0, 0xF4, 0xF8)
HR_COLOR       = RGBColor(0xCC, 0xCC, 0xCC)


def ensure_styles(doc):
    """Create custom styles if they don't exist."""
    styles = doc.styles

    def _ensure(name, base_style, font_size, bold=False, color=None,
                space_before=None, space_after=None, italic=False,
                font_name=None):
        try:
            return styles[name]
        except KeyError:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = styles[base_style] if base_style else None
            fmt = style.font
            fmt.size = Pt(font_size)
            fmt.bold = bold
            if color:
                fmt.color.rgb = color
            if italic:
                fmt.italic = True
            if font_name:
                fmt.name = font_name
            pf = style.paragraph_format
            if space_before is not None:
                pf.space_before = Pt(space_before)
            if space_after is not None:
                pf.space_after = Pt(space_after)
            return style

    # Heading styles — use built-in Word heading styles
    for level in range(1, 5):
        try:
            h = styles[f'Heading {level}']
        except KeyError:
            continue
        h.font.color.rgb = HEADING_COLOR
        if level == 1:
            h.font.size = Pt(24)
            h.font.bold = True
            h.paragraph_format.space_before = Pt(24)
            h.paragraph_format.space_after = Pt(12)
        elif level == 2:
            h.font.size = Pt(18)
            h.font.bold = True
            h.paragraph_format.space_before = Pt(18)
            h.paragraph_format.space_after = Pt(8)
        elif level == 3:
            h.font.size = Pt(14)
            h.font.bold = True
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(6)
        elif level == 4:
            h.font.size = Pt(12)
            h.font.bold = True
            h.paragraph_format.space_before = Pt(8)
            h.paragraph_format.space_after = Pt(4)

    # Code block style
    _ensure('Code Block', 'Normal', 9, font_name='Consolas',
            color=CODE_FG, space_before=2, space_after=2)

    # Blockquote style
    _ensure('Block Quote', 'Normal', 10, italic=True,
            color=RGBColor(0x55, 0x55, 0x55), space_before=4, space_after=4)

    # Subtitle style for version/date lines
    _ensure('Doc Subtitle', 'Normal', 11, color=RGBColor(0x66, 0x66, 0x66),
            space_before=2, space_after=2)

    return styles


def apply_inline_formatting(paragraph, text, default_bold=False):
    """Parse inline markdown (bold, code, links) and add formatted runs."""
    # Pattern to match **bold**, `code`, [text](url), or plain text
    pattern = re.compile(
        r'(\*\*(.+?)\*\*)'       # group 1,2: bold
        r'|(`(.+?)`)'             # group 3,4: inline code
        r'|(\[(.+?)\]\((.+?)\))' # group 5,6,7: link
        r'|([^*`\[]+)'           # group 8: plain text
        r'|(.)'                   # group 9: single special char
    )

    for m in pattern.finditer(text):
        if m.group(2):  # **bold**
            run = paragraph.add_run(m.group(2))
            run.bold = True
        elif m.group(4):  # `code`
            run = paragraph.add_run(m.group(4))
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        elif m.group(6):  # [text](url)
            run = paragraph.add_run(m.group(6))
            run.font.color.rgb = BRAND_BLUE
            run.underline = True
            # Add hyperlink as tooltip-like text
            # (Full hyperlink support requires oxml manipulation)
        elif m.group(8):  # plain
            run = paragraph.add_run(m.group(8))
            if default_bold:
                run.bold = True
        elif m.group(9):  # single char
            run = paragraph.add_run(m.group(9))
            if default_bold:
                run.bold = True


def add_table(doc, rows_text, insert_before=None):
    """Build a Word table from markdown table rows."""
    # Parse rows
    parsed = []
    for row_text in rows_text:
        cells = [c.strip() for c in row_text.strip('|').split('|')]
        # Skip separator rows (|---|---|)
        if all(re.match(r'^[-:]+$', c) for c in cells):
            continue
        parsed.append(cells)

    if len(parsed) < 1:
        return

    n_cols = max(len(row) for row in parsed)
    # Pad rows to same column count
    for row in parsed:
        while len(row) < n_cols:
            row.append('')

    table = doc.add_table(rows=len(parsed), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'

    for r_idx, row_data in enumerate(parsed):
        row = table.rows[r_idx]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ''  # clear default
            p = cell.paragraphs[0]
            apply_inline_formatting(p, cell_text)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)

            # Style header row
            if r_idx == 0:
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.size = Pt(9)
                # Blue background for header
                shading = parse_xml(
                    f'<w:shd {nsdecls("w")} w:fill="006CBF"/>')
                cell._element.get_or_add_tcPr().append(shading)
            else:
                for run in p.runs:
                    run.font.size = Pt(9)
                # Alternating row colors
                if r_idx % 2 == 0:
                    shading = parse_xml(
                        f'<w:shd {nsdecls("w")} w:fill="F2F6FA"/>')
                    cell._element.get_or_add_tcPr().append(shading)

    return table


def add_horizontal_rule(doc):
    """Add a thin horizontal line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # Add bottom border to simulate HR
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p


def add_blockquote(doc, text):
    """Add a blockquote with left border styling."""
    p = doc.add_paragraph()
    p.style = doc.styles['Block Quote']
    # Left indent
    p.paragraph_format.left_indent = Inches(0.4)
    # Add left border
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="12" w:space="8" w:color="006CBF"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    # Add shading
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F4F8" w:val="clear"/>')
    pPr.append(shading)
    apply_inline_formatting(p, text)
    return p


def add_code_block(doc, lines):
    """Add a code block with monospace font and background."""
    for line in lines:
        p = doc.add_paragraph()
        p.style = doc.styles['Code Block']
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = CODE_FG
        # Background shading
        pPr = p._element.get_or_add_pPr()
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
        pPr.append(shading)


def convert_document(input_path, output_path):
    """Main conversion: read markdown .docx, write formatted .docx."""
    doc = Document(input_path)
    new_doc = Document()

    # Set default font
    style = new_doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(6)

    # Page margins
    for section in new_doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    ensure_styles(new_doc)

    # Collect all paragraph text
    paragraphs = [p.text for p in doc.paragraphs]

    i = 0
    while i < len(paragraphs):
        text = paragraphs[i]
        stripped = text.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # ── Horizontal rule ──────────────────────────────────────
        if stripped in ('---', '***', '___'):
            add_horizontal_rule(new_doc)
            i += 1
            continue

        # ── Headings ─────────────────────────────────────────────
        heading_match = re.match(r'^(#{1,4})\s+(.+)$', text)
        if heading_match:
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2).strip()
            p = new_doc.add_paragraph()
            p.style = new_doc.styles[f'Heading {level}']
            apply_inline_formatting(p, heading_text)
            i += 1
            continue

        # ── Code fence ───────────────────────────────────────────
        if stripped.startswith('```'):
            code_lines = []
            i += 1
            while i < len(paragraphs) and not paragraphs[i].strip().startswith('```'):
                code_lines.append(paragraphs[i])
                i += 1
            add_code_block(new_doc, code_lines)
            i += 1  # skip closing ```
            continue

        # ── Table ────────────────────────────────────────────────
        if stripped.startswith('|') and '|' in stripped[1:]:
            table_rows = []
            while i < len(paragraphs) and paragraphs[i].strip().startswith('|'):
                table_rows.append(paragraphs[i])
                i += 1
            add_table(new_doc, table_rows)
            continue

        # ── Blockquote ───────────────────────────────────────────
        if stripped.startswith('> '):
            quote_text = stripped[2:]
            # Collect continuation lines
            i += 1
            while i < len(paragraphs) and paragraphs[i].strip().startswith('> '):
                quote_text += ' ' + paragraphs[i].strip()[2:]
                i += 1
            add_blockquote(new_doc, quote_text)
            continue

        # ── Bullet list ──────────────────────────────────────────
        if re.match(r'^(\s*)[-*]\s', text):
            indent = len(text) - len(text.lstrip())
            bullet_text = re.sub(r'^(\s*)[-*]\s+', '', text)
            p = new_doc.add_paragraph()
            p.style = new_doc.styles['List Bullet']
            if indent >= 2:
                try:
                    p.style = new_doc.styles['List Bullet 2']
                except KeyError:
                    p.style = new_doc.styles['List Bullet']
                p.paragraph_format.left_indent = Inches(0.75)
            apply_inline_formatting(p, bullet_text)
            i += 1
            continue

        # ── Numbered list ────────────────────────────────────────
        num_match = re.match(r'^(\d+)\.\s+(.+)$', text)
        if num_match:
            p = new_doc.add_paragraph()
            p.style = new_doc.styles['List Number']
            apply_inline_formatting(p, num_match.group(2))
            i += 1
            continue

        # ── Normal paragraph ─────────────────────────────────────
        p = new_doc.add_paragraph()
        apply_inline_formatting(p, text)
        i += 1

    new_doc.save(output_path)
    print(f"Saved: {output_path}")


if __name__ == '__main__':
    if len(sys.argv) != 3:
        print(f"Usage: {sys.argv[0]} input.docx output.docx")
        sys.exit(1)
    convert_document(sys.argv[1], sys.argv[2])
